import os
import sys 

from docx2pdf import convert

from docx import Document
from docx.shared import Pt
from docx.text.run import Run
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

code_dir_name = "code"
word_dir_name = "word"
pdf_dir_name = "pdf"

def formatFont(run: Run) -> None:

    font = run.font 
    font.name = "Aptos (Body)"
    font.size = Pt(13)


def create_docx (f_name: str, content: str) -> None: 
    
    if(content.count("\x1b") > 0):
        content = content.replace("\x1b", "\u241b")
    
    document = Document()

    code_paragraph = document.add_paragraph( )

    title = "//" + " " + f_name + "\n\n"
    title_run = code_paragraph.add_run(title)
    title_run.bold = True
    formatFont(title_run)

    code_run = code_paragraph.add_run(content)
    formatFont(code_run)

    # Add #pg in header
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_run = header_paragraph.add_run()

    qn('w:fldSimple')
    header_run._r.append(parse_xml(r'<w:fldSimple %s w:instr=" PAGE   \* MERGEFORMAT "><w:t>1</w:t></w:fldSimple>' % nsdecls('w')))

    # word file
    base_name ,_ = os.path.splitext(f_name) #tuple ('index.d', '.ts')
    docx_name = base_name + ".docx"
    docx_path = os.path.join(word_dir_name, docx_name)

    document.save(docx_path)

def check_dirs(*dir_names :str) -> None:
    for dir in dir_names:

        if not os.path.exists(dir):
            raise FileNotFoundError(f"Directory does not exist: {dir}")

if __name__ == "__main__":
    
    check_dirs( code_dir_name, word_dir_name, pdf_dir_name)

    for file_name in os.listdir(code_dir_name):
        print(f"File name: {file_name}")

        r_path = os.path.join(code_dir_name, file_name)

        with open(r_path, 'r', encoding='utf-8') as file:

            content = file.read()
            
            create_docx(file_name, content)

    convertTo: str = sys.argv[1]
    if(convertTo == "pdf"):
        convert(word_dir_name + "/", pdf_dir_name + "/")  



